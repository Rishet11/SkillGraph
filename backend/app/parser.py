from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re

from docx import Document
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    text: str
    warnings: list[str]
    source: str
    filename: str | None = None
    sections: dict[str, str] | None = None
    extraction_method: str = "text"


SECTION_ALIASES: dict[str, tuple[str, ...]] = {
    "summary": ("summary", "professional summary", "profile", "objective"),
    "skills": ("skills", "technical skills", "core skills", "competencies"),
    "experience": (
        "experience",
        "work experience",
        "professional experience",
        "employment",
    ),
    "projects": ("projects", "project experience", "selected projects"),
    "education": ("education", "academic background"),
    "certifications": ("certifications", "licenses", "certificates"),
    "required": (
        "required",
        "requirements",
        "must have",
        "minimum qualifications",
        "required qualifications",
    ),
    "preferred": (
        "preferred",
        "nice to have",
        "good to have",
        "bonus",
        "preferred qualifications",
    ),
    "responsibilities": (
        "responsibilities",
        "what you will do",
        "role responsibilities",
        "job responsibilities",
    ),
}


def _normalize_heading(line: str) -> str:
    stripped = line.strip().strip(":").lower()
    stripped = re.sub(r"\s+", " ", stripped)
    return stripped


def _is_heading_candidate(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if len(stripped) > 80:
        return False
    if stripped[0] in {"-", "*", "\u2022"}:
        return False
    if stripped.endswith("."):
        return False
    return True


def detect_sections(text: str) -> dict[str, str]:
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    section_lines: dict[str, list[str]] = {"unstructured": []}
    current_section = "unstructured"

    for line in lines:
        normalized = _normalize_heading(line)
        matched_section = None
        if _is_heading_candidate(line):
            for section, aliases in SECTION_ALIASES.items():
                if normalized in aliases:
                    matched_section = section
                    break
        if matched_section:
            current_section = matched_section
            section_lines.setdefault(current_section, [])
            continue
        section_lines.setdefault(current_section, []).append(line)

    cleaned_sections = {
        section: "\n".join(items).strip()
        for section, items in section_lines.items()
        if items
    }
    return cleaned_sections or {"unstructured": text}


def normalize_text(text: str) -> str:
    normalized_newlines = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in normalized_newlines.split("\n")]
    cleaned = "\n".join(line for line in lines if line)
    return cleaned.strip()


def parse_text_payload(text: str, source: str) -> ParsedDocument:
    normalized = normalize_text(text)
    warnings = []
    if len(normalized) < 60:
        warnings.append(f"{source} text is very short; analysis confidence may be lower.")
    sections = detect_sections(normalized)
    if list(sections) == ["unstructured"]:
        warnings.append(f"{source} section detection was limited; using unstructured parsing.")
    return ParsedDocument(
        text=normalized,
        warnings=warnings,
        source=source,
        sections=sections,
        extraction_method="text",
    )


def parse_uploaded_file(filename: str, content: bytes, source: str) -> ParsedDocument:
    suffix = Path(filename).suffix.lower()
    warnings: list[str] = []
    if suffix == ".txt":
        text = content.decode("utf-8", errors="ignore")
        normalized = normalize_text(text)
        sections = detect_sections(normalized)
        if list(sections) == ["unstructured"]:
            warnings.append("TXT section detection was limited; using unstructured parsing.")
        return ParsedDocument(
            text=normalized,
            warnings=warnings,
            source=source,
            filename=filename,
            sections=sections,
            extraction_method="txt",
        )
    if suffix == ".pdf":
        try:
            reader = PdfReader(BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = normalize_text("\n".join(pages))
            if not text:
                warnings.append("PDF parsing returned limited text. Paste the content for the best result.")
            sections = detect_sections(text)
            if list(sections) == ["unstructured"]:
                warnings.append("PDF section detection was limited; using unstructured parsing.")
            return ParsedDocument(
                text=text,
                warnings=warnings,
                source=source,
                filename=filename,
                sections=sections,
                extraction_method="pdf",
            )
        except Exception as exc:  # pragma: no cover - defensive
            raise ValueError(
                "Unable to parse PDF. Use paste-text fallback or a TXT/DOCX file."
            ) from exc
    if suffix == ".docx":
        try:
            document = Document(BytesIO(content))
            text = normalize_text("\n".join(p.text for p in document.paragraphs))
            sections = detect_sections(text)
            if list(sections) == ["unstructured"]:
                warnings.append("DOCX section detection was limited; using unstructured parsing.")
            return ParsedDocument(
                text=text,
                warnings=warnings,
                source=source,
                filename=filename,
                sections=sections,
                extraction_method="docx",
            )
        except Exception as exc:  # pragma: no cover - defensive
            raise ValueError(
                "Unable to parse DOCX. Use paste-text fallback or a TXT file."
            ) from exc
    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT, or paste text directly.")
